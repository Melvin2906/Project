from google import genai
from google.genai import types
import io
from datetime import datetime
from zoneinfo import ZoneInfo
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd
class GenAIExecption(Exception):
    """GenAI Exception base class"""

class ChatBot:
    CHATBOT_NAME = "My Gemini AI"
    MODEL_NAME = "gemini-2.5-flash"
    IMAGE_MODEL_NAME = "gemini-2.5-flash-image"

    BASE_INSTRUCTION = (
        "Please format your responses in clear Markdown with headings, lists, "
        "and emphasis when useful."
    )

    def __init__(self, api_key):
        self.client = genai.Client(api_key=api_key)
        self.conversation = None
        self._conversation_history = []
        self.preload_conversation()

    def _system_instruction(self, context=None):
        """`context` (ex: date/heure courante) part en system_instruction —
        jamais mélangé au message utilisateur. Avant, on le collait devant le
        prompt et le modèle le traitait comme un message à "accuser réception",
        d'où les réponses parasites ("Thank you for providing the system
        information...")."""
        if context:
            return f"{self.BASE_INSTRUCTION}\n\n{context}"
        return self.BASE_INSTRUCTION

    def _grounded_config(self, temperature, context=None):
        """google_search active le grounding : le modèle peut chercher sur le
        web avant de répondre, pour des infos à jour plutôt que sa seule
        mémoire d'entraînement."""
        return types.GenerateContentConfig(
            temperature=temperature,
            system_instruction=self._system_instruction(context),
            tools=[types.Tool(google_search=types.GoogleSearch())],
        )

    def _to_gemini_contents(self, messages):
        contents = []
        for m in messages:
            text = m.get("content")
            if text is None:
                text = m.get("parts", [""])[0]
            role = "model" if m["role"] == "assistant" else "user"
            contents.append(types.Content(role=role, parts=[types.Part(text=text)]))
        return contents

    def send_prompt(self, prompt, temperature=0.1, context=None):
        if temperature < 0 or temperature > 1:
            raise GenAIExecption('Temperature must be between 0 and 1')
        if not prompt:
            raise GenAIExecption('Prompt cannot be empty')

        try:
            response = self.client.models.generate_content(
                model=self.MODEL_NAME,
                contents=prompt,
                config=self._grounded_config(temperature, context),
            )
            return response.text.strip()
        except Exception as e:
            return f"Erreur Gemini: {e}"

    def send_prompt_with_history(self, messages, temperature=0.1, context=None):
        """messages : [{'role':.., 'content':..}, ...] reconstruit depuis la DB
        pour cette conversation précise — jamais d'état partagé entre
        utilisateurs. `context` va en system_instruction (voir plus haut)."""
        if not messages:
            raise GenAIExecption('Messages cannot be empty')

        contents = self._to_gemini_contents(messages)

        try:
            response = self.client.models.generate_content(
                model=self.MODEL_NAME,
                contents=contents,
                config=self._grounded_config(temperature, context),
            )
            return response.text.strip()
        except Exception as e:
            return f"Erreur Gemini: {e}"

    def send_prompt_with_image(self, prompt, image_bytes, temperature=0.1, context=None):
        if not prompt:
            raise GenAIExecption("Prompt cannot be empty")

        try:
            image_part = types.Part.from_bytes(data=image_bytes, mime_type="image/jpeg")
            response = self.client.models.generate_content(
                model=self.MODEL_NAME,
                contents=[prompt, image_part],
                config=types.GenerateContentConfig(
                    temperature=temperature,
                    system_instruction=self._system_instruction(context),
                ),
            )
            return response.text.strip()
        except Exception as e:
            return f"Erreur Gemini (image): {e}"

    def update_datetime(self, timezone):
        try:
            now = datetime.now(tz=ZoneInfo(timezone))
        except Exception:
            now = datetime.utcnow()
            timezone = "UTC"

        formatted = now.strftime("%Y-%m-%d %H:%M:%S")
        return (
            f"Current date and time: {formatted} ({timezone}). "
            f"Use this only if relevant to answering the user's question — "
            f"never mention or acknowledge receiving this information."
        )

    @property
    def history(self):
        return []

    def clear_conversation(self):
        pass
    def start_convertion(self):
        pass
    def preload_conversation(self, conversation_history=None):
        self._conversation_history = []

    def generate_image(self, prompt):
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")
        try:
            response = self.client.models.generate_content(
                model=self.IMAGE_MODEL_NAME,
                contents=[prompt],
            )
            for part in response.candidates[0].content.parts:
                if getattr(part, "inline_data", None) is not None:
                    return part.inline_data.data
            return None
        except Exception as e:
            print(f"Erreur de génération d'image: {e}")
            return None

    def generate_document(self, prompt, file_type):
        content_prompt = f"Génère le contenu textuel pour un fichier {file_type} basé sur : {prompt}. Sois concis."
        text_content = self.send_prompt(content_prompt)

        buf = io.BytesIO()

        if file_type == "pdf":
            p = canvas.Canvas(buf)
            p.drawString(100, 750, f"Document généré par {self.CHATBOT_NAME}")
            p.drawString(100, 730, text_content)
            p.showPage()
            p.save()
        elif file_type == "docx":
            doc = Document()
            doc.add_heading('Document IA', 0)
            doc.add_paragraph(text_content)
            doc.save(buf)
        elif file_type == "xlsx":
            df = pd.DataFrame({"Contenu": [text_content]})
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)

        buf.seek(0)
        return buf.getvalue()

    def read_document(self, prompt, file_bytes, mime_type, temperature=0.1, history=None):
        """`history` (optionnel) : messages précédents de la conversation
        (depuis la DB), pour ne pas dépendre d'un état interne partagé."""
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")

        try:
            doc_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)

            contents = []
            if history:
                contents.extend(self._to_gemini_contents(history))
            contents.append(types.Content(role="user", parts=[types.Part(text=prompt), doc_part]))

            response = self.client.models.generate_content(
                model=self.MODEL_NAME,
                contents=contents,
                config=types.GenerateContentConfig(temperature=temperature),
            )
            return response.text.strip()
        except Exception as e:
            return f"Erreur lors de l'analyse du document : {e}"
