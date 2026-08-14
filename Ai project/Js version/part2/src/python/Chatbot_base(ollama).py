import io
import os
import base64
import requests
from datetime import datetime
from zoneinfo import ZoneInfo
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd
from pypdf import PdfReader


class GenAIExecption(Exception):
    """GenAI Exception base class"""


class ChatBot:
    """Chat can only have one candidate count"""
    CHATBOT_NAME = "My Local AI"
    OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
    OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.1")
    OLLAMA_VISION_MODEL = os.environ.get("OLLAMA_VISION_MODEL", "llava")
    SD_MODEL_ID = os.environ.get("SD_MODEL_ID", "runwayml/stable-diffusion-v1-5")

    def __init__(self, api_key=None):
        self._api_key = api_key

        self.conversation = []
        self._conversation_history = []
        self._sd_pipeline = None
        self.preload_conversation()

    def _ollama_chat(self, messages, temperature=0.1):
        try:
            resp = requests.post(
                f"{self.OLLAMA_HOST}/api/chat",
                json={
                    "model": self.OLLAMA_MODEL,
                    "messages": messages,
                    "stream": False,
                    "options": {"temperature": temperature},
                },
                timeout=120,
            )
            resp.raise_for_status()
            data = resp.json()
            return data["message"]["content"].strip()
        except Exception as e:
            raise GenAIExecption(f"Erreur Ollama: {e}")

    def send_prompt(self, prompt, temperature=0.1):
        if temperature < 0 or temperature > 1:
            raise GenAIExecption('Temperature must be between 0 and 1')

        if not prompt:
            raise GenAIExecption('Prompt cannot be empty')

        self.conversation.append({"role": "user", "content": prompt})

        try:
            reply = self._ollama_chat(self.conversation, temperature)
            self.conversation.append({"role": "assistant", "content": reply})
            return reply
        except Exception as e:
            return f"Erreur LLM local: {e}"

    def send_prompt_with_history(self, messages, temperature=0.1, context=None):
        """Comme send_prompt, mais sans dépendre de l'état interne partagé
        (self.conversation). `messages` est la liste complète à envoyer
        ([{'role':..,'content':..}, ...]) — utile quand l'historique vit en DB
        par conversation/utilisateur, pour éviter que toutes les conversations
        de tous les utilisateurs ne se mélangent dans une seule instance globale.
        `context` (ex: date/heure courante) part en message "system" séparé,
        jamais mélangé au texte du message utilisateur."""
        if not messages:
            raise GenAIExecption('Messages cannot be empty')

        full_messages = list(self._conversation_history)
        if context:
            full_messages.append({"role": "system", "content": context})
        full_messages += list(messages)

        try:
            return self._ollama_chat(full_messages, temperature)
        except Exception as e:
            return f"Erreur LLM local: {e}"

    def send_prompt_with_image(self, prompt, image_bytes, temperature=0.1, context=None):
        if not prompt:
            raise GenAIExecption("Prompt cannot be empty")

        try:
            image_b64 = base64.b64encode(image_bytes).decode("utf-8")
            messages = []
            if context:
                messages.append({"role": "system", "content": context})
            messages.append({"role": "user", "content": prompt, "images": [image_b64]})

            resp = requests.post(
                f"{self.OLLAMA_HOST}/api/chat",
                json={
                    "model": self.OLLAMA_VISION_MODEL,
                    "messages": messages,
                    "stream": False,
                    "options": {"temperature": temperature},
                },
                timeout=120,
            )
            resp.raise_for_status()
            return resp.json()["message"]["content"].strip()

        except Exception as e:
            return f"Erreur LLM local (image): {e}"

    def update_datetime(self, timezone):
        try:
            now = datetime.now(tz=ZoneInfo(timezone))
        except Exception:
            now = datetime.utcnow()
            timezone = "UTC"

        formatted = now.strftime("%Y-%m-%d %H:%M:%S")

        return (
            f"System information:\n"
            f"- Current date and time: {formatted}\n"
            f"- Timezone: {timezone}\n"
            f"This information is authoritative."
        )

    @property
    def history(self):
        return [
            {"role": message["role"], "text": message["content"]}
            for message in self.conversation
        ]

    def clear_conversation(self):
        self.conversation = []

    def start_convertion(self):
        self.conversation = list(self._conversation_history)

    def _construct_message(self, text, role='user'):
        return {"role": role, "content": text}

    def preload_conversation(self, conversation_history=None):
        if isinstance(conversation_history, list):
            self._conversation_history = conversation_history
        else:
            self._conversation_history = [
                self._construct_message(
                    "Please format your responses in clear Markdown with headings, lists, and emphasis when useful."
                )
            ]

    def _load_sd_pipeline(self):
        if self._sd_pipeline is None:
            import torch
            from diffusers import StableDiffusionPipeline

            device = "cuda" if torch.cuda.is_available() else "cpu"
            dtype = torch.float16 if device == "cuda" else torch.float32

            self._sd_pipeline = StableDiffusionPipeline.from_pretrained(
                self.SD_MODEL_ID, torch_dtype=dtype
            ).to(device)

        return self._sd_pipeline

    def generate_image(self, prompt):
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")
        try:
            pipeline = self._load_sd_pipeline()
            image = pipeline(prompt).images[0]

            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            return img_byte_arr.getvalue()

        except Exception as e:
            print(f"Erreur de génération d'image: {e}")
            return None
    def generate_document(self, prompt, file_type):
        content_prompt = f"Génère le contenu textuel pour un fichier {file_type} basé sur : {prompt}. Sois concis."
        text_content = self.send_prompt(content_prompt)

        buf = io.BytesIO()

        if file_type == "pdf":
            p = canvas.Canvas(buf)
            p.drawString(100, 750, f"Document généré par {self.CHATBOT_NAME}")
            p.drawString(100, 730, text_content)
            p.showPage()
            p.save()

        elif file_type == "docx":
            doc = Document()
            doc.add_heading('Document IA', 0)
            doc.add_paragraph(text_content)
            doc.save(buf)

        elif file_type == "xlsx":
            df = pd.DataFrame({"Contenu": [text_content]})
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)

        buf.seek(0)
        return buf.getvalue()

    def _extract_text(self, file_bytes, mime_type):
        if mime_type == "application/pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)

        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""

    def read_document(self, prompt, file_bytes, mime_type, temperature=0.1, history=None):
        """Extrait le texte du document localement, puis l'envoie au LLM local en texte.
        `history` (optionnel) : messages précédents de la conversation (depuis la DB),
        pour ne pas dépendre de l'état interne partagé."""
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")

        try:
            extracted_text = self._extract_text(file_bytes, mime_type)

            if not extracted_text.strip():
                return "Impossible d'extraire du texte de ce document (format non supporté ou document scanné sans OCR)."

            full_prompt = f"{prompt}\n\n--- Contenu du document ---\n{extracted_text}"
            messages = list(self._conversation_history) + (history or []) + [
                {"role": "user", "content": full_prompt}
            ]
            return self.send_prompt_with_history(messages, temperature)

        except Exception as e:
            return f"Erreur lors de l'analyse du document : {e}"
