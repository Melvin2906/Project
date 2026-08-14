import google.generativeai as genai
from PIL import Image
import io
from datetime import datetime
from zoneinfo import ZoneInfo
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd

class GenAIExecption(Exception):
    """GenAI Exception base class"""

class ChatBot:
    """Chat can only have one candidate count"""
    CHATBOT_NAME = "My Gemini AI"

    def __init__(self, api_key):
        self.genai = genai
        self.genai.configure(api_key=api_key)
        self.model = self.genai.GenerativeModel("gemini-2.5-flash")
        self.conversation = None
        self._conversation_history = []

        self.preload_conversation()

    def send_prompt(self, prompt, temperature=0.1):
        if temperature < 0 or temperature > 1:
            raise GenAIExecption('Temperature must be between 0 and 1')

        if not prompt:
            raise GenAIExecption('Prompt cannot be empty')

        try:
            responce = self.conversation.send_message(
                content=prompt,
                generation_config=self._generation_config(temperature),
            )
            responce.resolve()
            return responce.text.strip()
        except Exception as e:
            return f"Erreur Gemini: {e}"

    def send_prompt_with_history(self, messages, temperature=0.1):
        """Comme send_prompt, mais reconstruit une conversation Gemini jetable à
        partir d'un historique explicite (mélange possible de deux formats :
        {'role':..,'content':..} venant de la DB, ou {'role':..,'parts':[..]}
        natif Gemini venant de self._conversation_history), au lieu de dépendre
        de self.conversation qui est partagé par toutes les requêtes de tous les
        utilisateurs sur cette instance de ChatBot. Indispensable dès qu'il y a
        plusieurs comptes/conversations en parallèle, sinon leurs échanges se
        mélangent."""
        if not messages:
            raise GenAIExecption('Messages cannot be empty')

        def _text_of(message):
            if "content" in message:
                return message["content"]
            return message["parts"][0]

        gemini_history = [
            {
                "role": "model" if m["role"] == "assistant" else m["role"],
                "parts": [_text_of(m)],
            }
            for m in messages[:-1]
        ]
        last_message = _text_of(messages[-1])

        try:
            temp_conversation = self.model.start_chat(history=gemini_history)
            response = temp_conversation.send_message(
                content=last_message,
                generation_config=self._generation_config(temperature),
            )
            response.resolve()
            return response.text.strip()
        except Exception as e:
            return f"Erreur Gemini: {e}"

    def send_prompt_with_image(self, prompt, image_bytes, temperature=0.1):
        if not prompt:
            raise GenAIExecption("Prompt cannot be empty")

        try:
            image = Image.open(io.BytesIO(image_bytes))

            response = self.model.generate_content(
                [
                    prompt,
                    image
                ],
                generation_config=self._generation_config(temperature)
            )

            response.resolve()
            return response.text.strip()

        except Exception as e:
            return f"Erreur Gemini (image): {e}"

    def update_datetime(self, timezone):
        try:
            now = datetime.now(tz=ZoneInfo(timezone))
        except Exception:
            now = datetime.utcnow()
            timezone = "UTC"

        formatted = now.strftime("%Y-%m-%d %H:%M:%S")

        return (
            f"System information:\n"
            f"- Current date and time: {formatted}\n"
            f"- Timezone: {timezone}\n"
            f"This information is authoritative."
        )

    @property
    def history(self):
        conversation_history = [
            {'role': message.role, 'text': message.parts[0].text} for message in self.conversation.history
        ]
        return conversation_history

    def clear_conversation(self):
        self.conversation = self.model.start_chat(history=[])

    def start_convertion(self):
        self.conversation = self.model.start_chat(history=self._conversation_history)

    def _generation_config(self, temperature):
        return genai.types.GenerationConfig(
            temperature=temperature
        )

    def _construct_message(self, text, role='user'):
        return {
            'role': role,
            'parts': [text]
        }

    def preload_conversation(self, conversation_history=None):
        if isinstance(conversation_history, list):
            self._conversation_history = conversation_history
        else:
            self._conversation_history = [
                self._construct_message("Please format your responses in clear Markdown with headings, lists, and emphasis when useful.")
            ]

    def generate_image(self, prompt):
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")
        try:
            imagen_model = self.genai.GenerativeModel("gemini-2.5-flash-image-preview")
            response = imagen_model.generate_content(prompt)
            generated_image = response.candidates[0].content.parts[0].inline_data.data
            img_byte_arr = io.BytesIO()
            generated_image.save(img_byte_arr, format='PNG')
            return img_byte_arr.getvalue()

        except Exception as e:
            print(f"Erreur de génération d'image: {e}")
            return None

    def generate_document(self, prompt, file_type):
        content_prompt = f"Génère le contenu textuel pour un fichier {file_type} basé sur : {prompt}. Sois concis."
        text_content = self.send_prompt(content_prompt)

        buf = io.BytesIO()

        if file_type == "pdf":
            p = canvas.Canvas(buf)
            p.drawString(100, 750, f"Document généré par {self.CHATBOT_NAME}")
            p.drawString(100, 730, text_content)
            p.showPage()
            p.save()

        elif file_type == "docx":
            doc = Document()
            doc.add_heading('Document IA', 0)
            doc.add_paragraph(text_content)
            doc.save(buf)

        elif file_type == "xlsx":
            df = pd.DataFrame({"Contenu": [text_content]})
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)

        buf.seek(0)
        return buf.getvalue()

    def read_document(self, prompt, file_bytes, mime_type, temperature=0.1, history=None):
        """Permet à Gemini d'analyser un document (PDF, Text, etc.).
        `history` (optionnel) : messages précédents de la conversation (depuis
        la DB), pour ne pas dépendre de l'état interne partagé self.conversation."""
        if not prompt:
            raise GenAIExecption("Le prompt ne peut pas être vide")

        try:
            document_data = {
                "mime_type": mime_type,
                "data": file_bytes
            }

            if history:
                gemini_history = [
                    {
                        "role": "model" if m["role"] == "assistant" else m["role"],
                        "parts": [m["content"]],
                    }
                    for m in history
                ]
                temp_conversation = self.model.start_chat(history=gemini_history)
                response = temp_conversation.send_message(
                    content=[prompt, document_data],
                    generation_config=self._generation_config(temperature),
                )
            else:
                response = self.model.generate_content(
                    [prompt, document_data],
                    generation_config=self._generation_config(temperature)
                )

            response.resolve()
            return response.text.strip()
        except Exception as e:
            return f"Erreur lors de l'analyse du document : {e}"
